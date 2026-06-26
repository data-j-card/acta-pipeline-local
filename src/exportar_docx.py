"""
Exporta el acta a .docx con el formato de un acta de referencia.

Une: encabezado (datos del usuario) + orden del día + antecedente + desarrollo
(de acta_borrador.txt) + tabla de asistencia (padron.csv) + cierre y
firmas (plantilla, con espacios en blanco para que el humano complete).

Uso:
    python src/exportar_docx.py
Salida: output/Acta.docx
"""
import csv
import re

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

import config
import encabezado
from pipeline import ORDEN_DEL_DIA

BORRADOR = config.OUTPUT_DIR / "acta_borrador.txt"
PADRON = config.OUTPUT_DIR / "padron.csv"
OUT = config.OUTPUT_DIR / "Acta.docx"

BLANCO = "________________________"


# ---------- helpers ----------
def _campo(parrafo, instr):
    """Inserta un campo de Word (p.ej. PAGE / NUMPAGES) en un parrafo."""
    run = parrafo.add_run()
    begin = OxmlElement("w:fldChar"); begin.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve"); instrText.text = instr
    end = OxmlElement("w:fldChar"); end.set(qn("w:fldCharType"), "end")
    run._r.append(begin); run._r.append(instrText); run._r.append(end)


def _footer_pagina(section):
    p = section.footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Página ")
    _campo(p, "PAGE")
    p.add_run(" de ")
    _campo(p, "NUMPAGES")


def _sombrear(cell, hexcolor):
    sh = OxmlElement("w:shd")
    sh.set(qn("w:val"), "clear")
    sh.set(qn("w:fill"), hexcolor)
    cell._tc.get_or_add_tcPr().append(sh)


def _seccion_titulo(doc, texto):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(texto)
    r.bold = True
    return p


def _parrafos_de(texto):
    return [b.strip() for b in re.split(r"\n\s*\n", texto.strip()) if b.strip()]


def _extraer(borrador_txt, ini, fin):
    seg = borrador_txt.split(ini, 1)[1]
    if fin:
        seg = seg.split(fin, 1)[0]
    return seg.strip()


# ---------- construccion ----------
def main():
    datos = encabezado.load(config.DATOS_FILE)
    titulo_apertura = encabezado.render_encabezado(datos)
    titulo, apertura = titulo_apertura.split("\n\n", 1)

    borrador = BORRADOR.read_text(encoding="utf-8")
    antecedente = _extraer(borrador, "ANTECEDENTE", "DESARROLLO DEL ORDEN DEL DÍA")
    desarrollo = _extraer(borrador,
                          "DESARROLLO DEL ORDEN DEL DÍA (segunda cita)", "[CIERRE")

    doc = Document()
    # estilo base
    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(11)
    for s in doc.sections:
        s.top_margin = s.bottom_margin = Cm(2)
        s.left_margin = s.right_margin = Cm(2.2)

    # --- encabezado de pagina (titulo, se repite) ---
    hp = doc.sections[0].header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    hr = hp.add_run(titulo.replace("\n", " — "))
    hr.bold = True
    hr.font.size = Pt(10)
    _footer_pagina(doc.sections[0])

    # --- apertura ---
    doc.add_paragraph(apertura)

    # --- orden del dia ---
    for i, item in enumerate(ORDEN_DEL_DIA, 1):
        doc.add_paragraph(f"{i}. {item}")

    # --- antecedente ---
    doc.add_paragraph()
    _seccion_titulo(doc, "ANTECEDENTE")
    for par in _parrafos_de(antecedente):
        doc.add_paragraph(par)

    # --- desarrollo ---
    doc.add_paragraph()
    _seccion_titulo(doc, "DESARROLLO DEL ORDEN DEL DÍA")
    for par in _parrafos_de(desarrollo):
        doc.add_paragraph(par)

    # --- tabla de asistencia / quorum ---
    doc.add_paragraph()
    _seccion_titulo(doc, "VERIFICACIÓN DEL QUÓRUM Y LISTADO DE ASISTENCIA")
    cols = ["APTO", "PROPIETARIO", "COEFICIENTE", "PODER", "ASISTIÓ", "COEF. QUÓRUM"]
    with open(PADRON, encoding="utf-8-sig") as f:
        filas = list(csv.DictReader(f))

    tbl = doc.add_table(rows=1, cols=len(cols))
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, c in enumerate(cols):
        cell = tbl.rows[0].cells[j]
        cell.text = c
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(8)
        _sombrear(cell, "D9D9D9")
    for fila in filas:
        celdas = tbl.add_row().cells
        vals = [fila["APTO"], fila["PROPIETARIO"], fila["COEFICIENTE"],
                fila["PODER"], fila["ASISTIO"], fila["COEF_QUORUM"]]
        for j, v in enumerate(vals):
            celdas[j].text = v
            celdas[j].paragraphs[0].runs and setattr(
                celdas[j].paragraphs[0].runs[0].font, "size", Pt(8))

    # --- cierre + firmas (plantilla del acta de referencia, con espacios) ---
    doc.add_paragraph()
    doc.add_paragraph(
        "Una vez agotado el orden del día y siendo las __________ se da por "
        "finalizada la Asamblea General Extraordinaria. En constancia firman:")
    doc.add_paragraph()
    _firmas(doc, "Presidente de la Asamblea", "Secretario(a) de la Asamblea")

    doc.add_paragraph()
    doc.add_paragraph(
        f"CONSTANCIA DE APROBACIÓN DEL ACTA N° {datos.get('numero_acta','XX')} "
        "DE ASAMBLEA EXTRAORDINARIA DE COPROPIETARIOS")
    doc.add_paragraph(
        "Los suscritos miembros de la comisión nombrada para aprobar el acta de esta "
        "Asamblea, una vez leída detenidamente y revisadas las proposiciones presentadas "
        "y las decisiones aprobadas, le impartimos nuestra aprobación unánime y se remite "
        f"a {datos.get('convoca','la administradora')} para los trámites pertinentes "
        "relativos a su publicación y para su registro en el libro de actas.")
    doc.add_paragraph()
    _firmas(doc, "Comisión verificadora del acta", "Comisión verificadora del acta")

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(f"LISTO -> {OUT}  ({len(filas)} filas en la tabla de asistencia)")


def _firmas(doc, cargo_izq, cargo_der):
    """Dos bloques de firma lado a lado con espacios en blanco."""
    p1 = doc.add_paragraph("____________________________          "
                           "____________________________")
    p2 = doc.add_paragraph(f"{BLANCO}          {BLANCO}")
    p3 = doc.add_paragraph("Apto ______                            Apto ______")
    p4 = doc.add_paragraph(f"{cargo_izq}".ljust(38) + cargo_der)
    for p in (p1, p2, p3, p4):
        for r in p.runs:
            r.font.size = Pt(10)


if __name__ == "__main__":
    main()
